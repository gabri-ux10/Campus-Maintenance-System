from __future__ import annotations

import datetime as dt
import os
import re
from collections import Counter
from pathlib import Path
from typing import Iterable

import javalang
from docx import Document
from fpdf import FPDF
from fpdf.errors import FPDFException


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "documentation" / "generated"
MARKDOWN_OUTPUT = OUTPUT_DIR / "Campus_Maintenance_System_Full_Documentation.md"
DOCX_OUTPUT = OUTPUT_DIR / "Campus_Maintenance_System_Full_Documentation.docx"
PDF_OUTPUT = OUTPUT_DIR / "Campus_Maintenance_System_Full_Documentation.pdf"

CODE_EXTENSIONS = {".java", ".js", ".jsx", ".ts", ".tsx", ".cpp", ".c", ".h", ".hpp", ".sql"}
EXCLUDED_DIRS = {
    ".git",
    ".idea",
    ".vscode",
    "node_modules",
    "dist",
    "build",
    "target",
    "__pycache__",
}


def read_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="ignore")


def collect_code_files(root: Path) -> list[Path]:
    files: list[Path] = []
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = [d for d in dirnames if d not in EXCLUDED_DIRS]
        for filename in filenames:
            path = Path(dirpath) / filename
            if path.suffix.lower() in CODE_EXTENSIONS:
                files.append(path)
    return sorted(files)


def compact_whitespace(value: str) -> str:
    return re.sub(r"\s+", " ", value.strip())


def parse_exports_js(content: str) -> list[str]:
    exports: set[str] = set()

    for match in re.finditer(r"(?m)^\s*export\s+default\s+([A-Za-z_$][\w$]*)", content):
        exports.add(f"default {match.group(1)}")
    for match in re.finditer(
        r"(?m)^\s*export\s+(?:async\s+)?function\s+([A-Za-z_$][\w$]*)\s*\(",
        content,
    ):
        exports.add(match.group(1))
    for match in re.finditer(
        r"(?m)^\s*export\s+(?:const|let|var|class)\s+([A-Za-z_$][\w$]*)",
        content,
    ):
        exports.add(match.group(1))
    for match in re.finditer(r"(?m)^\s*export\s*\{([^}]+)\}", content):
        names = [part.strip() for part in match.group(1).split(",") if part.strip()]
        for name in names:
            if " as " in name:
                exports.add(name.split(" as ", 1)[1].strip())
            else:
                exports.add(name)

    return sorted(exports)


def parse_js_like(content: str) -> dict:
    classes: set[str] = set()
    functions: set[str] = set()

    for match in re.finditer(r"(?m)^\s*(?:export\s+)?class\s+([A-Za-z_$][\w$]*)", content):
        classes.add(match.group(1))

    for match in re.finditer(
        r"(?m)^\s*(?:export\s+)?(?:async\s+)?function\s+([A-Za-z_$][\w$]*)\s*\(([^)]*)\)",
        content,
    ):
        name = match.group(1)
        params = compact_whitespace(match.group(2))
        functions.add(f"{name}({params})")

    for match in re.finditer(
        r"(?m)^\s*(?:export\s+)?(?:const|let|var)\s+([A-Za-z_$][\w$]*)\s*=\s*(?:async\s*)?\(([^)]*)\)\s*=>",
        content,
    ):
        name = match.group(1)
        params = compact_whitespace(match.group(2))
        functions.add(f"{name}({params})")

    for match in re.finditer(
        r"(?m)^\s*(?:export\s+)?(?:const|let|var)\s+([A-Za-z_$][\w$]*)\s*=\s*(?:async\s*)?([A-Za-z_$][\w$]*)\s*=>",
        content,
    ):
        name = match.group(1)
        params = compact_whitespace(match.group(2))
        functions.add(f"{name}({params})")

    return {
        "exports": parse_exports_js(content),
        "classes": sorted(classes),
        "functions": sorted(functions),
    }


def java_type_to_str(type_node) -> str:
    if type_node is None:
        return "void"

    if isinstance(type_node, javalang.tree.BasicType):
        base = type_node.name
    elif isinstance(type_node, javalang.tree.ReferenceType):
        base = type_node.name or ""
        if getattr(type_node, "arguments", None):
            args = []
            for arg in type_node.arguments:
                if hasattr(arg, "type") and arg.type is not None:
                    args.append(java_type_to_str(arg.type))
                elif hasattr(arg, "pattern_type") and arg.pattern_type is not None:
                    args.append(java_type_to_str(arg.pattern_type))
                else:
                    args.append("?")
            base = f"{base}<{', '.join(args)}>"
        if getattr(type_node, "sub_type", None) is not None:
            base = f"{base}.{java_type_to_str(type_node.sub_type)}"
    else:
        base = str(type_node)

    dimensions = getattr(type_node, "dimensions", None)
    if dimensions:
        base += "[]" * len(dimensions)
    return base


def java_param_to_str(param) -> str:
    param_type = java_type_to_str(param.type)
    if getattr(param, "varargs", False):
        param_type += "..."
    return f"{param_type} {param.name}"


def java_modifiers_to_str(modifiers: set[str]) -> str:
    order = ["public", "protected", "private", "static", "final", "abstract", "synchronized", "native"]
    ordered = [name for name in order if name in modifiers]
    remaining = sorted(m for m in modifiers if m not in order)
    return " ".join(ordered + remaining).strip()


def parse_java_with_regex(content: str) -> dict:
    type_matches = re.findall(r"\b(class|interface|enum)\s+([A-Za-z_][\w]*)", content)
    types = sorted({f"{kind} {name}" for kind, name in type_matches})

    method_pattern = re.compile(
        r"(?m)^\s*(?:public|protected|private)?\s*(?:static\s+)?(?:final\s+)?(?:abstract\s+)?"
        r"(?:synchronized\s+)?(?:native\s+)?(?:<[^>]+>\s+)?([\w<>\[\],.? ]+?)\s+([A-Za-z_][\w]*)\s*"
        r"\(([^)]*)\)\s*(?:throws[^{;]+)?\s*[;{]"
    )
    methods = []
    for return_type, name, params in method_pattern.findall(content):
        if name in {"if", "for", "while", "switch", "catch", "return", "throw"}:
            continue
        methods.append(f"{compact_whitespace(return_type)} {name}({compact_whitespace(params)})")
    methods = sorted(set(methods))

    if types:
        type_entries = [{"name": type_name, "kind": "type", "constructors": [], "methods": methods} for type_name in types]
    else:
        type_entries = [{"name": "unknown", "kind": "type", "constructors": [], "methods": methods}]

    return {"package": "", "types": type_entries, "parse_notes": "Regex fallback parser used."}


def parse_java(content: str) -> dict:
    try:
        tree = javalang.parse.parse(content)
    except Exception:
        return parse_java_with_regex(content)

    package_name = tree.package.name if tree.package else ""
    type_entries = []
    for path, node in tree.filter(javalang.tree.TypeDeclaration):
        parent_types = [parent.name for parent in path if isinstance(parent, javalang.tree.TypeDeclaration)]
        full_name = ".".join(parent_types + [node.name]) if parent_types else node.name
        kind = node.__class__.__name__.replace("Declaration", "").lower()

        constructors: list[str] = []
        for constructor in getattr(node, "constructors", []) or []:
            modifiers = java_modifiers_to_str(constructor.modifiers)
            params = ", ".join(java_param_to_str(param) for param in constructor.parameters)
            prefix = f"{modifiers} " if modifiers else ""
            constructors.append(f"{prefix}{constructor.name}({params})")

        methods: list[str] = []
        for method in getattr(node, "methods", []) or []:
            modifiers = java_modifiers_to_str(method.modifiers)
            params = ", ".join(java_param_to_str(param) for param in method.parameters)
            return_type = java_type_to_str(method.return_type)
            prefix = f"{modifiers} " if modifiers else ""
            methods.append(f"{prefix}{return_type} {method.name}({params})")

        type_entries.append(
            {
                "name": full_name,
                "kind": kind,
                "constructors": sorted(set(constructors)),
                "methods": sorted(set(methods)),
            }
        )

    return {"package": package_name, "types": sorted(type_entries, key=lambda item: item["name"]), "parse_notes": ""}


def parse_cpp(content: str) -> dict:
    classes = sorted(set(re.findall(r"(?m)^\s*(?:class|struct)\s+([A-Za-z_]\w*)", content)))
    function_pattern = re.compile(
        r"(?m)^\s*(?:template\s*<[^>]+>\s*)?(?:inline\s+)?(?:static\s+)?(?:virtual\s+)?"
        r"(?:constexpr\s+)?(?:[\w:&<>\*\~]+\s+)+([A-Za-z_~]\w*(?:::\w+)*)\s*\(([^;{}]*)\)\s*"
        r"(?:const)?\s*(?:noexcept)?\s*(?:=\s*0)?\s*[;{]"
    )
    disallowed = {"if", "for", "while", "switch", "catch", "return", "delete", "new"}
    functions = []
    for name, params in function_pattern.findall(content):
        if name in disallowed:
            continue
        functions.append(f"{name}({compact_whitespace(params)})")
    return {"classes": classes, "functions": sorted(set(functions))}


def parse_sql(content: str) -> dict:
    tables = sorted(
        set(
            re.findall(
                r"(?is)\bcreate\s+table\s+(?:if\s+not\s+exists\s+)?([A-Za-z_][\w.]*)",
                content,
            )
        )
    )
    views = sorted(
        set(
            re.findall(
                r"(?is)\bcreate\s+view\s+(?:if\s+not\s+exists\s+)?([A-Za-z_][\w.]*)",
                content,
            )
        )
    )
    routines = sorted(
        set(
            re.findall(
                r"(?is)\bcreate\s+(?:or\s+replace\s+)?(?:function|procedure)\s+([A-Za-z_][\w.]*)",
                content,
            )
        )
    )
    return {"tables": tables, "views": views, "routines": routines}


def infer_purpose(path: Path) -> str:
    parts = {segment.lower() for segment in path.parts}
    filename = path.stem

    if "controller" in parts or filename.endswith("Controller"):
        return "API/controller logic"
    if "service" in parts or filename.endswith("Service"):
        return "Business/service logic"
    if "repository" in parts or filename.endswith("Repository"):
        return "Data access layer"
    if "entity" in parts or filename.endswith("Entity"):
        return "Domain/entity model"
    if "dto" in parts or filename.endswith(("Request", "Response")):
        return "Data transfer model"
    if "config" in parts:
        return "Application/configuration code"
    if "pages" in parts:
        return "Frontend route/page component"
    if "components" in parts:
        return "Reusable UI component"
    if "hooks" in parts:
        return "Frontend hook/state logic"
    if "services" in parts:
        return "Frontend API/client service"
    if "utils" in parts:
        return "Shared utility/helper logic"
    if path.suffix.lower() == ".sql":
        return "Database schema/seed script"
    if path.suffix.lower() in {".cpp", ".h", ".hpp", ".c"}:
        return "C/C++ optimization/native module"
    return "Project source code"


def analyze_file(path: Path) -> dict:
    rel = path.relative_to(ROOT).as_posix()
    content = read_text(path)
    suffix = path.suffix.lower()
    data = {
        "path": rel,
        "language": suffix.lstrip("."),
        "purpose": infer_purpose(path),
        "details": {},
    }

    if suffix == ".java":
        data["details"] = parse_java(content)
    elif suffix in {".js", ".jsx", ".ts", ".tsx"}:
        data["details"] = parse_js_like(content)
    elif suffix in {".cpp", ".c", ".h", ".hpp"}:
        data["details"] = parse_cpp(content)
    elif suffix == ".sql":
        data["details"] = parse_sql(content)

    return data


def format_file_section(file_data: dict) -> list[str]:
    lines = [f"### `{file_data['path']}`", f"- Language: `{file_data['language']}`", f"- Purpose: {file_data['purpose']}"]
    details = file_data["details"]
    language = file_data["language"]

    if language == "java":
        package = details.get("package", "")
        if package:
            lines.append(f"- Package: `{package}`")
        parse_notes = details.get("parse_notes", "")
        if parse_notes:
            lines.append(f"- Parse Notes: {parse_notes}")
        types = details.get("types", [])
        if not types:
            lines.append("- Methods: No class/interface declarations detected.")
        else:
            lines.append("- Types and Methods:")
            for type_entry in types:
                kind = type_entry.get("kind", "type")
                lines.append(f"  - `{kind} {type_entry['name']}`")
                constructors = type_entry.get("constructors", [])
                methods = type_entry.get("methods", [])
                if constructors:
                    lines.append("    - Constructors:")
                    for constructor in constructors:
                        lines.append(f"      - `{constructor}`")
                if methods:
                    lines.append("    - Methods:")
                    for method in methods:
                        lines.append(f"      - `{method}`")
                if not constructors and not methods:
                    lines.append("    - No methods detected.")

    elif language in {"js", "jsx", "ts", "tsx"}:
        exports = details.get("exports", [])
        classes = details.get("classes", [])
        functions = details.get("functions", [])

        lines.append("- Exports:")
        if exports:
            lines.extend([f"  - `{item}`" for item in exports])
        else:
            lines.append("  - No named/default exports detected.")

        lines.append("- Classes:")
        if classes:
            lines.extend([f"  - `{item}`" for item in classes])
        else:
            lines.append("  - No class declarations detected.")

        lines.append("- Functions/Components:")
        if functions:
            lines.extend([f"  - `{item}`" for item in functions])
        else:
            lines.append("  - No function declarations detected.")

    elif language in {"cpp", "c", "h", "hpp"}:
        classes = details.get("classes", [])
        functions = details.get("functions", [])
        lines.append("- Classes/Structs:")
        if classes:
            lines.extend([f"  - `{item}`" for item in classes])
        else:
            lines.append("  - None detected.")
        lines.append("- Functions:")
        if functions:
            lines.extend([f"  - `{item}`" for item in functions])
        else:
            lines.append("  - None detected.")

    elif language == "sql":
        tables = details.get("tables", [])
        views = details.get("views", [])
        routines = details.get("routines", [])
        lines.append("- Tables:")
        if tables:
            lines.extend([f"  - `{item}`" for item in tables])
        else:
            lines.append("  - None detected.")
        lines.append("- Views:")
        if views:
            lines.extend([f"  - `{item}`" for item in views])
        else:
            lines.append("  - None detected.")
        lines.append("- Functions/Procedures:")
        if routines:
            lines.extend([f"  - `{item}`" for item in routines])
        else:
            lines.append("  - None detected.")

    lines.append("")
    return lines


def build_markdown(file_docs: Iterable[dict], extension_counts: Counter, generated_at: dt.datetime) -> str:
    lines = [
        "# Campus Maintenance System - Full Project Documentation",
        "",
        f"Generated on: **{generated_at.strftime('%Y-%m-%d %H:%M:%S')}**",
        "",
        "## Scope",
        "- This document covers all detected project code files (`.java`, `.js`, `.jsx`, `.ts`, `.tsx`, `.cpp`, `.c`, `.h`, `.hpp`, `.sql`).",
        "- For each file, the document lists purpose and detected functions/methods/classes.",
        "",
        "## Project Modules",
        "- `backend/`: Spring Boot API, business logic, repositories, entities, security, schedulers.",
        "- `frontend/`: React + Vite web app, dashboards, hooks, services, utility modules.",
        "- `database/`: SQL schema and seed scripts.",
        "- `backend/src/main/java/com/smartcampus/maintenance/optimization/`: Java optimization routines used for scoring and safe image handling.",
        "- `tests/`: Integration and E2E flows.",
        "",
        "## Code Inventory",
    ]

    for ext, count in sorted(extension_counts.items()):
        lines.append(f"- `{ext}`: {count} file(s)")
    lines.extend(["", "## File-by-File Function Documentation", ""])

    for file_doc in file_docs:
        lines.extend(format_file_section(file_doc))

    return "\n".join(lines)


def markdown_to_docx(markdown_text: str, output_path: Path) -> None:
    document = Document()
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif not line:
            document.add_paragraph("")
        else:
            document.add_paragraph(line)
    document.save(output_path)


def markdown_to_pdf(markdown_text: str, output_path: Path) -> None:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()
    pdf.set_font("Helvetica", size=10)

    def wrap_long_line(line: str, max_chars: int = 120) -> list[str]:
        if len(line) <= max_chars:
            return [line]

        segments: list[str] = []
        current = []
        length = 0
        for token in line.split(" "):
            if len(token) > max_chars:
                if current:
                    segments.append(" ".join(current))
                    current = []
                    length = 0
                for i in range(0, len(token), max_chars):
                    segments.append(token[i : i + max_chars])
                continue

            projected = len(token) if not current else length + 1 + len(token)
            if projected > max_chars:
                segments.append(" ".join(current))
                current = [token]
                length = len(token)
            else:
                current.append(token)
                length = projected

        if current:
            segments.append(" ".join(current))
        return segments

    def write_line(line: str, line_height: int = 5) -> None:
        for segment in wrap_long_line(line):
            sanitized = segment.encode("latin-1", errors="replace").decode("latin-1")
            if not sanitized.strip():
                pdf.ln(line_height)
                continue
            pdf.set_x(pdf.l_margin)
            try:
                pdf.multi_cell(0, line_height, text=sanitized)
            except FPDFException:
                # Fallback for edge-case tokens that cannot be wrapped by the engine.
                for i in range(0, len(sanitized), 80):
                    chunk = sanitized[i : i + 80]
                    if chunk:
                        pdf.set_x(pdf.l_margin)
                        pdf.multi_cell(0, line_height, text=chunk)

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if line.startswith("# "):
            pdf.set_font("Helvetica", style="B", size=16)
            write_line(line[2:].strip(), line_height=8)
            pdf.ln(1)
            pdf.set_font("Helvetica", size=10)
        elif line.startswith("## "):
            pdf.set_font("Helvetica", style="B", size=13)
            write_line(line[3:].strip(), line_height=7)
            pdf.ln(1)
            pdf.set_font("Helvetica", size=10)
        elif line.startswith("### "):
            pdf.set_font("Helvetica", style="B", size=11)
            write_line(line[4:].strip(), line_height=6)
            pdf.set_font("Helvetica", size=10)
        elif line.startswith("- "):
            write_line(f"* {line[2:].strip()}")
        elif not line:
            pdf.ln(2)
        else:
            write_line(line)

    pdf.output(str(output_path))


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    code_files = collect_code_files(ROOT)
    extension_counts = Counter(path.suffix.lower().lstrip(".") for path in code_files)
    file_docs = [analyze_file(path) for path in code_files]
    file_docs.sort(key=lambda item: item["path"])

    generated_at = dt.datetime.now()
    markdown_text = build_markdown(file_docs, extension_counts, generated_at)

    MARKDOWN_OUTPUT.write_text(markdown_text, encoding="utf-8")
    markdown_to_docx(markdown_text, DOCX_OUTPUT)
    markdown_to_pdf(markdown_text, PDF_OUTPUT)

    print(f"Generated markdown: {MARKDOWN_OUTPUT}")
    print(f"Generated Word document: {DOCX_OUTPUT}")
    print(f"Generated PDF document: {PDF_OUTPUT}")


if __name__ == "__main__":
    main()
